from azure.storage.blob import BlobServiceClient, BlobClient
from datetime import datetime
from docx import Document as DocxDocument
from dotenv import load_dotenv
from functools import wraps
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_aws import ChatBedrock
from langchain_core.prompts import PromptTemplate
from langchain_core.messages import HumanMessage
from langchain_core.output_parsers import JsonOutputParser
from langchain_openai import AzureChatOpenAI
from llama_parse import LlamaParse
from llama_index.readers.azstorage_blob import AzStorageBlobReader
from llama_index.core import Document as LlamaParseDocument
from openai import AzureOpenAI
from pdf2image import convert_from_bytes
from PIL import Image
from pydantic import BaseModel, Field
from typing import List, Optional, Any, Tuple, Dict, Union

import base64
import cv2
import httpx
import io
import logging
import nest_asyncio
import numpy as np
import os
import re
import tempfile
import tiktoken
import time


# Configure logging here
log_file_path = "app.log"
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[logging.FileHandler(log_file_path), logging.StreamHandler()]
)
logger = logging.getLogger(__name__)
def retry_on_exception(retries=3, backoff_factor=1, allowed_exceptions=(Exception,)):
    """
    Decorator to retry a function if it raises an exception with exponential backoff.
    """
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            attempt = 0
            while attempt < retries:
                try:
                    return func(*args, **kwargs)
                except allowed_exceptions as e:
                    attempt += 1
                    sleep_time = backoff_factor * attempt
                    logger.warning(f"Error in {func.__name__}: {e}. Retrying in {sleep_time}s (Attempt {attempt}/{retries}).")
                    time.sleep(sleep_time)
            return func(*args, **kwargs)  # Final attempt
        return wrapper
    return decorator

PARSER_OCR_VALIDATION_PROMPT = """Evaluate the supplied text and decide whether it contains substantive content (e.g., factual data, narrative, instructions, dialogue, etc.) rather than a refusal or error message.

What counts as a refusal/error message
Typical indicators include—but are not limited to—phrases such as:

“I’m sorry, but I can’t help with that.”

“Sorry, I can’t comply with that request.”

“I am not able to…” / “I cannot provide…”

“This violates policy…” / “forbidden request”

Generic error notices (“Error 403”, “Request denied”, etc.)

If the text primarily consists of apologies, disclaimers, policy statements, or error codes with no meaningful substantive content, treat it as a refusal/error.

Response format
Return only a lowercase Boolean:

true — the text contains substantive/valid content

false — the text is a refusal or error message

Do not include any additional words, punctuation, or explanation—just true or false.

{format_instructions}

text:
\"\"\"
{text}
\"\"\"
"""

OCR_PROMPT = """You are an expert linguist tasked with carefully analyzing and extracting text from an image. Your goal is to extract all visible text exactly as it appears, preserving formatting, without making any corrections or alterations.

Please follow these instructions:

1. Examine the image carefully, looking for any text in any language or any numbers.

2. Extract every character, word, and numeral you can identify, exactly as it appears. Do not correct spelling, grammar, or formatting—even if you believe there are errors.

3. Preserve all original formatting, including:

         Line breaks
         Tables and their structure
         Indentations
         Capitalization
         Punctuation
         Spacing between words and characters

4. If any character is unreadable or unclear, replace it with [?].

5. Take your time to inspect different fonts, sizes, and placements—sometimes numbers or letters may be embedded within graphics or between lines. Do not rush.

Please begin your analysis and provide the extracted text below without any additional commentary."""

class IsValidOCRText(BaseModel):
    """
    Validation of OCR text.
    """
    is_valid: bool = Field(..., description="Indicates if the OCR text is valid")

class CustomHTTPClient(httpx.Client):
    def __init__(self, *args, **kwargs):
        kwargs.pop("proxies", None)  # Remove the 'proxies' argument if present
        super().__init__(*args, **kwargs)

class PDFToBase64Converter:
    def __init__(
        self,
        dpi: int = 300,
        max_size_mb: float = 3.0,
        max_dimension: Optional[int] = None,
        brightness_threshold: int = 100,
        beta: int = 50,
        border_size: int = 150,
    ):
        """
        :param dpi: resolution for rendering each PDF page
        :param max_size_mb: target max JPEG size per page
        :param max_dimension: max width/height in px (None = no resize)
        :param brightness_threshold: threshold for brightness correction
        :param beta: amount to boost brightness when below threshold
        :param border_size: white border to add around deskewed text
        """
        self.dpi = dpi
        self.max_size_mb = max_size_mb
        self.max_dimension = max_dimension
        self.brightness_threshold = brightness_threshold
        self.beta = beta
        self.border_size = border_size

    def process_pdf_to_base64(self, pdf_path: str) -> List[str]:
        """
        Converts every page of the PDF into a preprocessed JPEG and returns
        a list of base64-encoded strings.
        """
        # 1) load PDF bytes
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

        # 2) render pages to PIL Images
        pil_pages = convert_from_bytes(pdf_bytes, dpi=self.dpi)

        result_b64: List[str] = []
        for pil in pil_pages:
            # 3a) PIL -> OpenCV (BGR)
            cv_img = cv2.cvtColor(np.array(pil), cv2.COLOR_RGB2BGR)

            # 3b) preprocess
            proc = self.preprocess_image(
                cv_img,
                brightness_threshold=self.brightness_threshold,
                beta=self.beta,
                border_size=self.border_size
            )

            # 3c) back to PIL (RGB)
            pil_proc = Image.fromarray(cv2.cvtColor(proc, cv2.COLOR_BGR2RGB))

            # 3d) save under size limit to a temp file
            tmp = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
            tmp_path = tmp.name
            tmp.close()
            self.save_image_under_size_limit(
                pil_proc,
                tmp_path,
                max_size_mb=self.max_size_mb,
                max_dimension=self.max_dimension
            )

            # 3e) read & encode
            with open(tmp_path, "rb") as imgf:
                jpg = imgf.read()
            os.remove(tmp_path)

            b64 = base64.b64encode(jpg).decode("utf-8")
            result_b64.append(b64)

        return result_b64

    def save_image_under_size_limit(
        self,
        pil_image: Image.Image,
        output_path: str,
        max_size_mb: float = 3.0,
        max_dimension: Optional[int] = None
    ) -> None:
        """
        Save a PIL image as JPEG, resizing and reducing quality until it's under
        max_size_mb and max_dimension.
        """
        # 1) convert RGBA -> RGB
        if pil_image.mode == "RGBA":
            pil_image = pil_image.convert("RGB")

        # 2) enforce dimension limit
        if max_dimension:
            w, h = pil_image.size
            if max(w, h) > max_dimension:
                scale = max_dimension / max(w, h)
                pil_image = pil_image.resize(
                    (int(w * scale), int(h * scale)),
                    resample=Image.Resampling.LANCZOS
                )

        # 3) iteratively reduce quality
        quality = 95
        max_bytes = int(max_size_mb * 1024 * 1024)
        while quality >= 10:
            buf = io.BytesIO()
            pil_image.save(buf, format="JPEG", quality=quality, optimize=True)
            data = buf.getvalue()
            if len(data) <= max_bytes:
                with open(output_path, "wb") as f:
                    f.write(data)
                return
            quality -= 5

        # final fallback at lowest quality
        with open(output_path, "wb") as f:
            pil_image.save(f, format="JPEG", quality=quality, optimize=True)

    def preprocess_image(
        self,
        image: np.ndarray,
        brightness_threshold: int,
        beta: int,
        border_size: int
    ) -> np.ndarray:
        """
        Full preprocess: brightness boost, binarize, denoise, deskew, crop/pad borders.
        """
        # brightness
        mean_b = self.compute_mean_brightness(image)
        if mean_b < brightness_threshold:
            image = self.increase_brightness(image, beta)

        # grayscale & Otsu binarization
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        _, bw = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        # noise removal
        denoised = self._noise_removal(bw)

        # deskew
        # deskewed = self._deskew(denoised)

        # crop to content & add white border
        # cropped = self._remove_borders(deskewed)
        padded = cv2.copyMakeBorder(
            denoised,
            top=border_size, bottom=border_size,
            left=border_size, right=border_size,
            borderType=cv2.BORDER_CONSTANT,
            value=[255, 255, 255]
        )

        # to BGR
        return cv2.cvtColor(padded, cv2.COLOR_GRAY2BGR)

    def compute_mean_brightness(self, image: np.ndarray) -> float:
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        return float(np.mean(gray))

    def increase_brightness(self, image: np.ndarray, beta: int) -> np.ndarray:
        return cv2.convertScaleAbs(image, alpha=1.0, beta=beta)

    def _noise_removal(self, img: np.ndarray) -> np.ndarray:
        kernel = np.ones((1, 1), np.uint8)
        m = cv2.dilate(img, kernel, iterations=1)
        m = cv2.erode(m, kernel, iterations=1)
        m = cv2.morphologyEx(m, cv2.MORPH_CLOSE, kernel)
        return cv2.medianBlur(m, 3)

    def _get_skew_angle(self, img: np.ndarray) -> float:
        blur = cv2.GaussianBlur(img, (9, 9), 0)
        th = cv2.threshold(blur, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (30, 5))
        dil = cv2.dilate(th, kernel, iterations=2)
        cnts, _ = cv2.findContours(dil, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)
        if not cnts:
            return 0.0
        rect = cv2.minAreaRect(max(cnts, key=cv2.contourArea))
        angle = rect[-1]
        return -angle if angle > -45 else -(90 + angle)

    def _rotate(self, img: np.ndarray, angle: float) -> np.ndarray:
        h, w = img.shape[:2]
        M = cv2.getRotationMatrix2D((w//2, h//2), angle, 1.0)
        return cv2.warpAffine(img, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)

    def _deskew(self, img: np.ndarray) -> np.ndarray:
        return self._rotate(img, self._get_skew_angle(img))

    def _remove_borders(self, img: np.ndarray) -> np.ndarray:
        cnts, _ = cv2.findContours(img, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        if not cnts:
            return img
        x, y, w, h = cv2.boundingRect(max(cnts, key=cv2.contourArea))
        return img[y:y+h, x:x+w]
    
class BaseLLMClass:
    def __init__(self) -> None:
        load_dotenv()
    
    def get_llm(self, 
                temperature: float = 0.7, 
                max_tokens: int = 1024, 
                provider: str = 'azure', 
                azure_deployment_model: str = 'gpt-4.1-mini',
                aws_model_name: str = 'claude-3.7-sonnet',
                timeout: float = 16) -> Any:
        
        # Mapping of AWS model names to their model IDs.
        aws_model_map = {
            'claude-3.7-sonnet': "us.anthropic.claude-3-7-sonnet-20250219-v1:0",
            'deepseek-r1': "us.deepseek.r1-v1:0"
            # Add additional mappings as needed.
        }
        
        if provider == 'aws':
            model_id = aws_model_map.get(aws_model_name)
            if not model_id:
                raise ValueError(f"Unknown AWS model: {aws_model_name}")
            return ChatBedrock(
                aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
                aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
                model=model_id,
                region="us-east-1"
            )
        
        openai_api_version = "2025-01-01-preview"
        
        if azure_deployment_model.startswith('o'):
            temperature = 1
            openai_api_version = "2025-01-01-preview"
            max_tokens = None
        
        # Default to AzureChatOpenAI
        return AzureChatOpenAI(
            openai_api_version=openai_api_version,
            azure_deployment=azure_deployment_model,
            max_tokens=max_tokens,
            temperature=temperature,
            timeout=timeout,
            http_client=CustomHTTPClient()
        )
    
    def text_splitter(self, chunk_size: int = 1000, chunk_overlap: int = 100, documents: str = None, metadata: Dict = None, is_separator_regex=False):
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            is_separator_regex=is_separator_regex
        )

        return text_splitter.create_documents([documents], [metadata])
    
    def split_text_into_chunks(self, text, chunk_size=70000, encoding_name='cl100k_base'):
        encoding = tiktoken.get_encoding(encoding_name)
        tokens = encoding.encode(text)
        chunks = [tokens[i:i + chunk_size] for i in range(0, len(tokens), chunk_size)]
        text_chunks = [encoding.decode(chunk) for chunk in chunks]
        return text_chunks
    
    def invoke_llm_with_fallback(
        self,
        prompt_template: Any,
        data: Dict[str, Any] | str,
        provider: str = "azure",
        temperature: float = 0.2,
        max_tokens: Optional[int] = None,
        azure_deployment_model: str = "o1-mini",
        aws_model_name: str = "claude-3.7-sonnet",
        timeout: float = 60,
        **kwargs
    ) -> Union[str, Dict[str, Any]]:
        """
        Invoke the LLM pipeline with a given prompt template and data.
        If the primary provider fails, it falls back to AWS.

        Args:
            prompt_template (Any): The prompt template for the LLM.
            data (dict): The input data dictionary to be passed to the LLM.
            provider (str): LLM provider ("azure", "aws", "google"). Default is "azure".
            temperature (float): Sampling temperature for the LLM.
            max_tokens (Optional[int]): Maximum tokens allowed for response.
            azure_deployment_model (str): Azure LLM model to use.
            aws_model_name (str): AWS LLM model to use.
            **kwargs: Additional keyword arguments for configuring the pipeline (e.g., JsonOutputParser).

        Returns:
            Union[str, Dict[str, Any]]: Processed output from the LLM.
        """
        try:
            # Get the primary LLM
            llm = self.get_llm(
                temperature=temperature, 
                max_tokens=max_tokens, 
                provider=provider, 
                azure_deployment_model=azure_deployment_model,
                timeout=timeout,
            )

            # Build the LLM pipeline
            llm_chain = prompt_template | llm
            for key, value in kwargs.items():
                llm_chain |= value  # Chain additional components dynamically

            chain_response = llm_chain.invoke(data)

        except Exception as error:
            logger.error(f"Primary provider ({provider}) failed: {error}. Falling back to AWS.")

            # Get AWS fallback LLM
            llm = self.get_llm(provider="aws", aws_model_name=aws_model_name)
            llm_chain = prompt_template | llm
            for key, value in kwargs.items():
                llm_chain |= value  # Apply same transformations in fallback

            chain_response = llm_chain.invoke(data)

        if isinstance(chain_response, dict):
            return chain_response
        
        final_chain_response = chain_response.content
        final_chain_response = re.sub(r'```markdown|```|# Patient Summary|# Clinical Notes Template', '', final_chain_response)
        return final_chain_response

class CustomAzStorageBlobReader(BaseLLMClass):
    """
    Wrapper around AzStorageBlobReader that:
      - Handles .pdf via OCR
      - Handles .txt/.json/.md via raw text
      - Handles .doc/.docx via python-docx
      - Falls back to AzStorageBlobReader.load_resource for everything else
    """
    def __init__(self, container_name: str, connection_string: str, blob: Optional[str] = None):
        super().__init__()
        nest_asyncio.apply()
        load_dotenv()

        self.container_name = container_name
        self.connection_string = connection_string
        self.blob = blob  # optional single‐blob mode

        self.reader = AzStorageBlobReader(
            container_name=container_name,
            connection_string=connection_string
        )
        # our own Azure + OCR clients
        self._svc = BlobServiceClient.from_connection_string(connection_string)
        self.claude_client    = ChatBedrock(
            model="us.anthropic.claude-3-7-sonnet-20250219-v1:0",
            region="us-east-1"
        )
        self.azure_o4_mini_client = AzureOpenAI(
            azure_deployment="o4-mini",
            api_version="2025-01-01-preview",
            timeout=300
        )
        self.azure_gpt_4o_client = AzureOpenAI(
            azure_deployment="gpt-4o",
            api_version="2025-01-01-preview",
            timeout=300
        )

        self.LLAMAINDEX_PARSER = LlamaParse(
            result_type="markdown",
            use_vendor_multimodal_model=True,
            azure_openai_deployment_name="gpt-4o",
            azure_openai_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            azure_openai_api_version="2025-01-01-preview",
            azure_openai_key=os.getenv("AZURE_OPENAI_API_KEY"),
        )

        # precompile skip regex for our special extensions
        self.SPECIAL_EXT_RE = re.compile(r"\.(pdf|txt|json|jpg|jpeg|md|docx?)$", re.IGNORECASE)
        self.pdf_converter = PDFToBase64Converter(
            dpi=300,
            max_size_mb=3.0,
            max_dimension=2000,
            brightness_threshold=100,
            beta=50,
            border_size=150,
        )

    @retry_on_exception(retries=3, backoff_factor=10, allowed_exceptions=(Exception,))
    def _handle_pdf(self, client: BlobClient, name: str) -> List[LlamaParseDocument]:
        raw = client.download_blob().readall()
        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        tmp.write(raw); tmp_path = tmp.name; tmp.close()
        try:
            b64_list = self.pdf_converter.process_pdf_to_base64(tmp_path)
        finally:
            os.remove(tmp_path)

        out: List[LlamaParseDocument] = []
        for idx, b64 in enumerate(b64_list, start=1):
            provider, text = self.get_ocr_text(b64)
            out.append(LlamaParseDocument(
                text=text,
                metadata={
                    "name":     name,
                    "page":     idx,
                    "provider": provider,
                    "fetched":  datetime.now().isoformat(),
                }
            ))
        return out

    @retry_on_exception(retries=3, backoff_factor=10, allowed_exceptions=(Exception,))
    def _handle_text(self, client: BlobClient, name: str) -> List[LlamaParseDocument]:
        raw = client.download_blob().readall()
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("latin-1", errors="ignore")

        out: List[LlamaParseDocument] = []
        for i in range(0, len(text), 10_000):
            out.append(LlamaParseDocument(
                text=text[i:i+10_000],
                metadata={"name": name, "fetched": datetime.now().isoformat()}
            ))
        return out

    @retry_on_exception(retries=3, backoff_factor=10, allowed_exceptions=(Exception,))
    def _handle_doc(self, client: BlobClient, name: str) -> List[LlamaParseDocument]:
        raw = client.download_blob().readall()
        stream = io.BytesIO(raw)
        docx = DocxDocument(stream)
        full = "\n".join(p.text for p in docx.paragraphs)

        out: List[LlamaParseDocument] = []
        for i in range(0, len(full), 10_000):
            out.append(LlamaParseDocument(
                text=full[i:i+10_000],
                metadata={"name": name, "fetched": datetime.now().isoformat()}
            ))
        return out

    @retry_on_exception(retries=3, backoff_factor=10, allowed_exceptions=(Exception,))
    def _handle_image(self, client: BlobClient, name: str) -> List[LlamaParseDocument]:
        raw = client.download_blob().readall()
        nparr = np.frombuffer(raw, np.uint8)
        cv_img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        if cv_img is None:
            return []

        proc = self.pdf_converter.preprocess_image(
            cv_img,
            brightness_threshold=self.pdf_converter.brightness_threshold,
            beta=self.pdf_converter.beta,
            border_size=self.pdf_converter.border_size
        )
        pil_img = Image.fromarray(cv2.cvtColor(proc, cv2.COLOR_BGR2RGB))

        tmp = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
        tmp_path = tmp.name; tmp.close()
        self.pdf_converter.save_image_under_size_limit(
            pil_img,
            tmp_path,
            max_size_mb=self.pdf_converter.max_size_mb,
            max_dimension=self.pdf_converter.max_dimension
        )
        with open(tmp_path, "rb") as f:
            jpg_data = f.read()
        os.remove(tmp_path)

        b64 = base64.b64encode(jpg_data).decode("utf-8")
        provider, text = self.get_ocr_text(b64)

        return [LlamaParseDocument(
            text=text,
            metadata={"name": name, "provider": provider, "fetched": datetime.now().isoformat()}
        )]

    def load_data(self) -> List[LlamaParseDocument]:
        container = self._svc.get_container_client(self.reader.container_name)
        docs: List[LlamaParseDocument] = []

        # decide which blobs to process
        if self.blob:
            container_client = self._svc.get_container_client(self.container_name)
            credential = self._svc.credential
            # if it's a URL, SDK grabs the blob_name for us
            if self.blob.startswith(("http://", "https://")):
                blob_client = BlobClient.from_blob_url(
                    self.blob,
                    credential=credential
                )
                # sanity-check container
                if blob_client.container_name != self.container_name:
                    raise ValueError(
                        f"URL container '{blob_client.container_name}' "
                        f"does not match expected '{self.container_name}'"
                    )
                name = blob_client.blob_name

            # otherwise treat it as a plain blob name
            else:
                blob_client = container_client.get_blob_client(self.blob)
                name = self.blob

            # —— existence check only in single-blob mode —— 
            if not blob_client.exists():
                raise FileNotFoundError(
                    f"Blob '{name}' does not exist in container '{self.container_name}'"
                )

            blob_names = [name]
        else:
            blob_names = [b.name for b in container.list_blobs()]

        for name in blob_names:
            ext    = os.path.splitext(name)[1].lower()
            client = container.get_blob_client(name)

            if self.SPECIAL_EXT_RE.search(name):
                if ext == ".pdf":
                    docs.extend(self._handle_pdf(client, name))
                elif ext in {".txt", ".json", ".md"}:
                    docs.extend(self._handle_text(client, name))
                elif ext in {".doc", ".docx"}:
                    docs.extend(self._handle_doc(client, name))
                elif ext in {".jpg", ".jpeg"}:
                    docs.extend(self._handle_image(client, name))
                else:
                    # should not hit
                    file_extractor = {ext: self.LLAMAINDEX_PARSER}
                    loader = self.reader

                    loader.blob = name
                    loader.file_extractor = file_extractor

                    docs.extend(loader.load_resource(name))
            else:
                # Use the existing LLAMAPARSE loader as default for Othe files.
                file_extractor = {ext: self.LLAMAINDEX_PARSER}
                loader = self.reader

                loader.blob = name
                loader.file_extractor = file_extractor

                docs.extend(loader.load_resource(name))

        return docs

    def get_ocr_text(self, base64_image: str) -> Tuple[str, str]:
        providers = [
            ("gpt-4o", self.call_gpt4o_ocr),
            ("o4-mini", self.call_o4mini_ocr),
            ("claude",  self.call_claude_ocr),
        ]
        for name, fn in providers:
            try:
                text = fn(base64_image)
                return name, text
            except Exception as e:
                logger.warning(f"OCR provider {name} failed: {e}")
        raise RuntimeError("All OCR providers failed")
    
    def is_valid_ocr_text(self, text: str) -> bool:
        template = PARSER_OCR_VALIDATION_PROMPT
        parser = JsonOutputParser(pydantic_object=IsValidOCRText)
        prompt_template = PromptTemplate(
            input_variables=["text"],
            partial_variables={"format_instructions": parser.get_format_instructions()},
            template=template,
        )

        resp = self.invoke_llm_with_fallback(
            prompt_template=prompt_template, 
            data={"text": text}, 
            provider="azure", 
            azure_deployment_model="gpt-4.1-mini",
            parser=parser
        )
        
        return resp.get("is_valid", False)

    @retry_on_exception(retries=3, backoff_factor=10, allowed_exceptions=(Exception,))
    def call_gpt4o_ocr(self, b64: str) -> str:
        messages = [
            {"role":"system","content":OCR_PROMPT},
            {"role":"user","content":[{"type":"image_url","image_url":{"url":f"data:image/jpeg;base64,{b64}"}}]},
            {"role":"user","content":"Please perform OCR on the above image."}
        ]
        resp = self.azure_gpt_4o_client.chat.completions.create(
            model="gpt-4o", messages=messages, temperature=0.1
        )
        text = resp.choices[0].message.content
        if not self.is_valid_ocr_text(text):
            raise Exception("Invalid OCR output from GPT-4o")
        return text

    @retry_on_exception(retries=3, backoff_factor=10, allowed_exceptions=(Exception,))
    def call_o4mini_ocr(self, b64: str) -> str:
        messages = [
            {"role":"system","content":OCR_PROMPT},
            {"role":"user","content":[{"type":"image_url","image_url":{"url":f"data:image/jpeg;base64,{b64}"}}]},
            {"role":"user","content":"Please perform OCR on the above image."}
        ]
        resp = self.azure_o4_mini_client.chat.completions.create(
            model="o4-mini", messages=messages
        )
        text = resp.choices[0].message.content
        if not self.is_valid_ocr_text(text):
            raise Exception("Invalid OCR output from o4-mini")
        return text

    @retry_on_exception(retries=3, backoff_factor=10, allowed_exceptions=(Exception,))
    def call_claude_ocr(self, b64: str) -> str:
        msg = HumanMessage(content=[
            {"type":"text","text":OCR_PROMPT},
            {"type":"image","source":{"type":"base64","media_type":"image/jpeg","data":b64}},
            {"type":"text","text":"Please perform OCR on the above image."}
        ])
        text = self.claude_client.invoke([msg]).content
        if not self.is_valid_ocr_text(text):
            raise Exception("Invalid OCR output from Claude")
        return text

loader = CustomAzStorageBlobReader(
    connection_string="",
    container_name="",
    blob=""
)

print(loader.load_data())  # Example usage to load data from the specified container
